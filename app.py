# ==============================================
# GERADOR DE DECLARACOES E CRONOGRAMA DE DEFESAS
# Versao Streamlit Cloud - UFF Instituto de Quimica
# CORRIGIDO PARA INCLUIR TODOS OS MEMBROS DA BANCA E HORARIO
# ==============================================

import streamlit as st
import pandas as pd
from datetime import datetime
import io
import zipfile
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== CONFIGURACAO DA PAGINA =====
st.set_page_config(
    page_title="Gerador de Declarações - UFF",
    page_icon="📄",
    layout="wide"
)

# ===== ESTILOS CSS =====
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1e3a5f;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .success-box {
        padding: 1rem;
        background-color: #d4edda;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
    }
    .stProgress > div > div > div > div {
        background-color: #1e3a5f;
    }
</style>
""", unsafe_allow_html=True)

# ===== INICIALIZAR STATE =====
if 'processado' not in st.session_state:
    st.session_state.processado = False
if 'zip_buffer' not in st.session_state:
    st.session_state.zip_buffer = None
if 'periodo_letivo' not in st.session_state:
    st.session_state.periodo_letivo = "2025.2"

# ===== FUNCOES DO SEU CODIGO ORIGINAL (COLAB) =====
def formatar_nome(nome):
    """Formata o nome para manter a capitalizacao correta"""
    if pd.isna(nome) or nome == "":
        return ""
    return str(nome).title()

def obter_grau_curso(curso):
    """Define o grau e curso com base no tipo de curso"""
    if pd.isna(curso) or curso == "":
        return "Bacharelado em Química"

    curso = str(curso).lower().strip()

    if 'industrial' in curso:
        return "Bacharelado em Química Industrial"
    elif 'licenciatura' in curso:
        return "Licenciatura em Química"
    else:
        return "Bacharelado em Química"

def formatar_horario(horario):
    """Formata o horario para pegar apenas a hora de inicio"""
    if pd.isna(horario) or horario == "":
        return "horário a definir"

    horario_str = str(horario)

    # Remove conteudo entre parenteses
    if '(' in horario_str and ')' in horario_str:
        horario_str = horario_str.split('(')[0].strip()

    # Pega apenas a primeira parte antes de qualquer separador
    separadores = ['/', 'às', '-', 'as', 'a']
    for sep in separadores:
        if sep in horario_str:
            horario_str = horario_str.split(sep)[0].strip()
            break

    # Remove espacos extras e garante o formato
    horario_str = horario_str.replace(' ', '')

    # Adiciona "horas" se nao tiver
    if 'h' not in horario_str.lower():
        horario_str += 'h'

    return horario_str

def formatar_data_sem_dia_semana(data_input):
    """Formata a data removendo o dia da semana entre parenteses"""
    if pd.isna(data_input) or data_input == "":
        return "data a definir"

    data_str = str(data_input)

    # Remove qualquer conteudo entre parenteses (dia da semana)
    if '(' in data_str and ')' in data_str:
        data_str = data_str.split('(')[0].strip()

    # Tenta converter para o formato padrao dd/mm/aaaa
    try:
        data_obj = pd.to_datetime(data_str, dayfirst=True, errors='coerce')
        if not pd.isna(data_obj):
            return data_obj.strftime('%d/%m/%Y')
        else:
            return data_str
    except:
        return data_str

def extrair_horario_local(horario_str):
    """Extrai horario e local de uma string combinada"""
    if pd.isna(horario_str) or horario_str == "":
        return "horário a definir", "local a definir"

    horario_str = str(horario_str)

    # Extrair local (se estiver entre parenteses)
    local = "Sala a definir"
    if '(' in horario_str and ')' in horario_str:
        match = re.search(r'\((.*?)\)', horario_str)
        if match:
            local = match.group(1)
        horario_str = horario_str.split('(')[0].strip()

    # Formatar horario
    horario = horario_str.strip()

    return horario, local

def configurar_paragrafo(document, texto, negrito=False, italico=False, tamanho=12, alinhamento='left'):
    """Configura um paragrafo com formatacao especifica"""
    p = document.add_paragraph()

    if alinhamento == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alinhamento == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(texto)
    run.font.size = Pt(tamanho)

    if negrito:
        run.bold = True
    if italico:
        run.italic = True

    return p

def set_cell_background(cell, color):
    """Define a cor de fundo de uma celula da tabela"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell_border(cell):
    """Define bordas para todas as celulas"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    # Definir todas as bordas
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)

    tcPr.append(tcBorders)

def criar_tabela_cronograma_unica(doc, dados_defesas, periodo_letivo):
    """Cria uma unica tabela continua com texto informativo e dados das defesas"""

    # Calcular numero total de linhas: 1 linha de texto informativo + 1 linha de cabecalho + linhas de dados
    total_linhas = 1 + 1 + len(dados_defesas)

    # Criar tabela unica com 6 colunas
    tabela = doc.add_table(rows=total_linhas, cols=6)
    tabela.style = 'Table Grid'
    tabela.autofit = False

    # Larguras das colunas (otimizadas para uso do espaco)
    widths = [Inches(1.4), Inches(1.0), Inches(0.8), Inches(0.8), Inches(1.3), Inches(2.2)]

    # Aplicar larguras as colunas
    for row in tabela.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # LINHA 1: Texto informativo (mesclando todas as 6 colunas)
    linha_info = tabela.rows[0]

    # Mesclar celulas da primeira linha corretamente
    cell_info = linha_info.cells[0].merge(linha_info.cells[5])

    # Limpar qualquer texto existente
    cell_info.text = ""

    # Adicionar texto informativo
    textos_info = [
        f"Informamos abaixo o cronograma das defesas de monografia referente ao período letivo {periodo_letivo}.",
        "",
        "Lembramos que os estudantes de nossos cursos que assistirem às defesas de monografia têm",
        "direito ao aproveitamento da respectiva carga horária como atividade complementar, bastando",
        "assinar a respectiva lista de presença.",
        "",
        "Obs.: As defesas serão realizadas no Anfiteatro ou em salas do Instituto de Química. O local",
        "segue na coluna \"Horário/ local\""
    ]

    for texto in textos_info:
        p = cell_info.add_paragraph()
        if texto != "":
            run = p.add_run(texto)
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # Aplicar bordas à celula de informacoes
    set_cell_border(cell_info)

    # LINHA 2: Cabecalho da tabela
    linha_cabecalho = tabela.rows[1]
    cabecalho_campos = ["Nome do aluno", "Data", "Horário", "Local", "Orientador", "Título do trabalho"]

    for i, campo in enumerate(cabecalho_campos):
        if i < len(linha_cabecalho.cells):
            cell = linha_cabecalho.cells[i]
            cell.text = campo

            # Formatar cabecalho em negrito e com fundo cinza
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            set_cell_background(cell, "D9D9D9")
            set_cell_border(cell)

    # LINHAS 3+: Dados das defesas
    for idx, dados in enumerate(dados_defesas):
        linha_idx = idx + 2
        if linha_idx < len(tabela.rows):
            linha_dados = tabela.rows[linha_idx]
        else:
            linha_dados = tabela.add_row()

        campos = [
            dados['nome_aluno'],
            dados['data'],
            dados['horario'],
            dados['local'],
            dados['orientador'],
            dados['titulo']
        ]

        for i, valor in enumerate(campos):
            if i < len(linha_dados.cells):
                cell = linha_dados.cells[i]
                cell.text = valor
                set_cell_border(cell)

                # Ajustar fonte dos dados
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    return tabela

def gerar_cronograma_defesas(df, periodo_letivo="2025.2"):
    """Gera o documento Word com o cronograma de defesas"""

    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(40)
        section.right_margin = Pt(40)

    # Titulo do cronograma (em negrito e centralizado)
    titulo = f"Cronograma de defesas de monografia {periodo_letivo}"
    configurar_paragrafo(doc, titulo, negrito=True, tamanho=14, alinhamento='center')

    doc.add_paragraph()  # Espaco

    # Preparar dados para a tabela
    dados_defesas = []

    # Processar todas as linhas da planilha
    start_index = 0
    if len(df) > 0 and 'Carimbo de data/hora' in str(df.iloc[0, 0]):
        start_index = 1

    for index in range(start_index, len(df)):
        linha = df.iloc[index]

        # Verificar se é uma linha válida
        if pd.notna(linha.get('Nome do aluno')) and str(linha['Nome do aluno']).strip() not in ['', 'Nome do aluno']:

            # Extrair e formatar dados
            nome_aluno = formatar_nome(linha['Nome do aluno'])

            # Extrair data para ordenacao
            try:
                data_input = str(linha['Escolha a data para a defesa'])

                # Converter para datetime para ordenacao (com dayfirst=True)
                if '(' in data_input and ')' in data_input:
                    data_sem_parenteses = data_input.split('(')[0].strip()
                else:
                    data_sem_parenteses = data_input

                data_ordenacao = pd.to_datetime(data_sem_parenteses, dayfirst=True, errors='coerce')

                # USAR A NOVA FUNCAO PARA FORMATAR A DATA SEM DIA DA SEMANA
                data_display = formatar_data_sem_dia_semana(data_input)

            except Exception as e:
                data_ordenacao = datetime.max
                data_display = formatar_data_sem_dia_semana(str(linha['Escolha a data para a defesa']))

            # Extrair horario e local
            horarios = []
            for col in linha.index:
                if 'horário' in str(col).lower() and pd.notna(linha[col]) and linha[col] != "":
                    horarios.append(str(linha[col]))
                elif 'horario' in str(col).lower() and pd.notna(linha[col]) and linha[col] != "":
                    horarios.append(str(linha[col]))

            if horarios:
                horario, local = extrair_horario_local(horarios[0])
            else:
                horario, local = "horário a definir", "local a definir"

            # Extrair horario para ordenacao
            try:
                # Tentar extrair hora para ordenacao cronologica
                horario_ordenacao = re.search(r'(\d{1,2})h', horario)
                if horario_ordenacao:
                    hora = int(horario_ordenacao.group(1))
                else:
                    hora = 0
            except:
                hora = 0

            # Outros dados
            orientador = formatar_nome(linha['Orientador'])
            titulo_trabalho = linha['Título da Defesa']

            dados_defesas.append({
                'nome_aluno': nome_aluno,
                'data': data_display,
                'horario': horario,
                'local': local,
                'orientador': orientador,
                'titulo': titulo_trabalho,
                'data_ordenacao': data_ordenacao if not pd.isna(data_ordenacao) else datetime.max,
                'hora_ordenacao': hora
            })

    # Ordenar por data e horario
    dados_defesas.sort(key=lambda x: (x['data_ordenacao'], x['hora_ordenacao']))

    # Criar tabela unica com informacoes e dados
    criar_tabela_cronograma_unica(doc, dados_defesas, periodo_letivo)

    # Adicionar linhas vazias no final (como no modelo)
    for _ in range(6):
        doc.add_paragraph()

    return doc

def gerar_documento_word(dados):
    """Gera um documento Word com a declaracao - VERSÃO COMPLETA DO COLAB"""

    # Formatar os dados
    nome_aluno = formatar_nome(dados['Nome do aluno'])
    matricula = str(dados['Matrícula'])
    titulo = dados["Título da Defesa"]
    curso_original = dados['Curso']
    curso_formatado = obter_grau_curso(curso_original)

    # Extrair data e formatar (remover dia da semana entre parenteses)
    try:
        data_input = str(dados['Escolha a data para a defesa'])
        # Remove qualquer conteudo entre parenteses
        if '(' in data_input and ')' in data_input:
            data_input = data_input.split('(')[0].strip()
        # Usar dayfirst=True para datas no formato brasileiro
        data_defesa = pd.to_datetime(data_input, dayfirst=True).strftime('%d/%m/%Y')
    except Exception as e:
        data_defesa = str(dados['Escolha a data para a defesa'])
        # Remove qualquer conteudo entre parenteses
        if '(' in data_defesa and ')' in data_defesa:
            data_defesa = data_defesa.split('(')[0].strip()

    # Extrair e formatar horario
    horarios = []
    for col in dados.index:
        if 'horário' in str(col).lower() and pd.notna(dados[col]) and dados[col] != "":
            horarios.append(str(dados[col]))
        elif 'horario' in str(col).lower() and pd.notna(dados[col]) and dados[col] != "":
            horarios.append(str(dados[col]))

    horario = formatar_horario(horarios[0] if horarios else "")

    # Membros da banca - BUSCANDO AS COLUNAS ESPECIFICAS
    orientador = formatar_nome(dados.get('Orientador', ''))
    membro_titular1 = formatar_nome(dados.get('Membro titular 1', ''))
    membro_titular2 = formatar_nome(dados.get('Membro Titular 2', ''))
    membro_suplente = formatar_nome(dados.get('Membro Suplente', ''))
    coorientador = formatar_nome(dados.get('Coorientador', ''))

    # Criar documento Word
    doc = Document()

    # Configurar margens (aproximadamente)
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(60)
        section.right_margin = Pt(60)

    # Cabecalho
    configurar_paragrafo(doc, "SERVIÇO PÚBLICO FEDERAL", tamanho=12, alinhamento='center')
    configurar_paragrafo(doc, "MINISTÉRIO DA EDUCAÇÃO", tamanho=12, alinhamento='center')
    configurar_paragrafo(doc, "UNIVERSIDADE FEDERAL FLUMINENSE", tamanho=12, alinhamento='center')
    configurar_paragrafo(doc, "PRÓ-REITORIA DE GRADUAÇÃO", tamanho=12, alinhamento='center')

    doc.add_paragraph()  # Espaco

    # Titulo DECLARACAO
    configurar_paragrafo(doc, "DECLARAÇÃO", negrito=True, tamanho=14, alinhamento='center')

    doc.add_paragraph()  # Espaco

    # Texto da declaracao em UM UNICO PARAGRAFO
    texto_declaracao = f"Declaro, para os devidos fins, que a Banca Examinadora da Monografia de Final de Curso do(a) estudante {nome_aluno.upper()}, matriculado(a) sob o número {matricula}, cujo título é \"{titulo}\" defendida e aprovada no dia {data_defesa}, às {horario}, para obtenção do grau {curso_formatado}, foi composta pelos seguintes membros: TITULARES: {orientador} (Orientador(a))"

    # Adicionar membros titulares
    if membro_titular1 and membro_titular1 != "":
        texto_declaracao += f", {membro_titular1}"

    if membro_titular2 and membro_titular2 != "":
        texto_declaracao += f", {membro_titular2}"

    # Adicionar suplentes
    texto_declaracao += f"; SUPLENTES: {membro_suplente}"

    # Adicionar coorientador como suplente se existir
    if coorientador and coorientador != "":
        texto_declaracao += f", {coorientador} (Coorientador(a))"

    texto_declaracao += "."

    p_declaracao = configurar_paragrafo(doc, texto_declaracao, tamanho=12)
    p_declaracao.paragraph_format.first_line_indent = Pt(0)

    # Espacos para o rodape
    for _ in range(12):
        doc.add_paragraph()

    # Rodape
    data_emissao = datetime.now().strftime('%d/%m/%Y às %H:%M:%S')
    configurar_paragrafo(doc, "Universidade Federal Fluminense", tamanho=12, alinhamento='center')
    configurar_paragrafo(doc, f"Niterói, {data_emissao}", tamanho=12, alinhamento='center')

    doc.add_paragraph()  # Espaco

    configurar_paragrafo(doc, "CPF [CPF_DO_COORDENADOR]", tamanho=10, alinhamento='center')

    doc.add_paragraph()  # Espaco

    texto_rodape1 = "Este documento foi gerado pelo Sistema Acadêmico da Universidade Federal Fluminense - IdUFF."
    texto_rodape2 = 'Para verificar a autenticidade deste documento, acesse https://inscricao.id.uff.br no link "Validar declaração"'

    configurar_paragrafo(doc, texto_rodape1, tamanho=9, alinhamento='center')
    configurar_paragrafo(doc, texto_rodape2, tamanho=9, alinhamento='center')

    doc.add_paragraph()  # Espaco
    configurar_paragrafo(doc, "[CÓDIGO_DE_VALIDAÇÃO]", tamanho=9, alinhamento='center')

    return doc

# ===== INTERFACE PRINCIPAL =====
st.markdown('<p class="main-header">Gerador de Declarações e Cronograma</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Universidade Federal Fluminense - Instituto de Química</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Desenvolvido por  <strong>Tadeu L. Araujo</strong> - versão 1.0</p>', unsafe_allow_html=True)

# Sidebar com instrucoes
with st.sidebar:
    st.header("Instruções")
    st.markdown("""
    1. Clique no botão "Browse files" e faça upload do arquivo Excel com os dados das defesas
    2. Configure o período letivo
    3. Clique em "Processar"
    4. Baixe o arquivo ZIP com todos os documentos
    
    **Formato esperado:**
    - Arquivo Excel (.xlsx)
    - Aba: "Respostas ao formulário 1"
    - Colunas necessárias:
        - Nome do aluno
        - Matrícula
        - Curso
        - Título da Defesa
        - Orientador
        - Membro titular 1
        - Membro Suplente
        - Membro Titular 2 (opcional)
        - Coorientador (opcional)
        - Escolha a data para a defesa
        - Coluna com "horário" no nome
    """)
    
    st.markdown("---")
    st.info("**Conteúdo do ZIP:**")
    st.markdown("""
    - Declarações individuais (todos os membros da banca)
    - Cronograma completo de defesas
    """)

# Upload do arquivo
uploaded_file = st.file_uploader(
    "Selecione o arquivo Excel com os dados das defesas",
    type=['xlsx', 'xls'],
    help="Arquivo Excel exportado do Google Forms"
)

# Configuracoes
periodo_letivo = st.text_input("Período Letivo", value="2026.1")
st.session_state.periodo_letivo = periodo_letivo

# Detectar período do nome do arquivo
if uploaded_file is not None:
    padroes_periodo = [r'(\d{4}\.\d)', r'(\d{4}-\d)', r'(\d{4}_\d)']
    for padrao in padroes_periodo:
        match = re.search(padrao, uploaded_file.name)
        if match:
            periodo_detectado = match.group(1).replace('-', '.').replace('_', '.')
            st.info(f"Período detectado no nome do arquivo: {periodo_detectado}")
            periodo_letivo = periodo_detectado
            st.session_state.periodo_letivo = periodo_letivo
            break

# Botão de processamento
if uploaded_file and st.button("🔄 Processar Arquivo e Gerar ZIP", type="primary", use_container_width=True):
    
    with st.spinner("Processando arquivo..."):
        try:
            # Ler o arquivo Excel
            try:
                df = pd.read_excel(uploaded_file, sheet_name='Respostas ao formulário 1')
            except:
                try:
                    df = pd.read_excel(uploaded_file, sheet_name='Respostas ao formulario 1')
                except:
                    df = pd.read_excel(uploaded_file)
            
            st.success(f"✅ Planilha carregada com {len(df)} registros")
            
            # Mostrar preview
            with st.expander("📊 Visualizar dados e colunas"):
                st.dataframe(df.head())
                st.write("**Colunas encontradas:**")
                st.write(list(df.columns))
            
            # Criar ZIP em memoria
            zip_buffer = io.BytesIO()
            contador_declaracoes = 0
            
            # Barra de progresso
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Determinar inicio dos dados
            start_index = 0
            if len(df) > 0 and 'Carimbo de data/hora' in str(df.iloc[0, 0]):
                start_index = 1
                st.info("Pulando linha de cabeçalho...")
            
            total_registros = len(df) - start_index
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # Gerar declarações individuais
                if total_registros > 0:
                    status_text.text("Gerando declarações individuais...")
                    
                    for index in range(start_index, len(df)):
                        linha = df.iloc[index]
                        
                        # Verificar se é uma linha válida
                        if pd.notna(linha.get('Nome do aluno')) and str(linha['Nome do aluno']).strip() not in ['', 'Nome do aluno']:
                            
                            # Gerar documento Word da declaração
                            try:
                                doc = gerar_documento_word(linha)
                                
                                # Nome do arquivo
                                nome_aluno_arquivo = str(linha['Nome do aluno']).replace(' ', '_').replace('/', '_').replace('\\', '_')
                                nome_arquivo_word = f"Declaracao_{nome_aluno_arquivo}.docx"
                                
                                # Salvar no buffer
                                doc_buffer = io.BytesIO()
                                doc.save(doc_buffer)
                                doc_buffer.seek(0)
                                
                                # Adicionar ao ZIP
                                zipf.writestr(nome_arquivo_word, doc_buffer.getvalue())
                                contador_declaracoes += 1
                                
                            except Exception as e:
                                st.warning(f"Erro ao gerar declaração para {linha['Nome do aluno']}: {e}")
                        
                        # Atualizar progresso
                        progress_bar.progress((index - start_index + 1) / total_registros)
                    
                    status_text.text(f"✅ {contador_declaracoes} declarações geradas")
                
                # Gerar cronograma
                status_text.text("Gerando cronograma de defesas...")
                progress_bar.progress(0.9)
                
                try:
                    doc_cronograma = gerar_cronograma_defesas(df, periodo_letivo)
                    
                    # Nome do arquivo do cronograma
                    nome_arquivo_cronograma = f"Cronograma_defesas_monografia_{periodo_letivo}.docx"
                    
                    # Salvar no buffer
                    cronograma_buffer = io.BytesIO()
                    doc_cronograma.save(cronograma_buffer)
                    cronograma_buffer.seek(0)
                    
                    # Adicionar ao ZIP
                    zipf.writestr(nome_arquivo_cronograma, cronograma_buffer.getvalue())
                    
                    st.success(f"✅ Cronograma gerado: {nome_arquivo_cronograma}")
                    
                except Exception as e:
                    st.error(f"❌ Erro ao gerar cronograma: {e}")
                
                progress_bar.progress(1.0)
                status_text.text("✅ Processamento concluído!")
            
            # Preparar ZIP para download
            zip_buffer.seek(0)
            st.session_state.zip_buffer = zip_buffer
            st.session_state.processado = True
            
            st.balloons()
            
            # Resumo
            st.success(f"""
            ## ✅ Processamento Concluído!
            
            **Resultados:**
            - 📄 **{contador_declaracoes}** declarações individuais geradas
            - 📅 **1** cronograma de defesas gerado
            - 📦 Arquivo ZIP pronto para download
            
            **Período letivo:** {periodo_letivo}
            """)
            
        except Exception as e:
            st.error(f"❌ Erro ao processar arquivo: {e}")
            import traceback
            st.code(traceback.format_exc())

# Botão de download do ZIP (após processamento)
if st.session_state.processado and st.session_state.zip_buffer:
    st.markdown("---")
    st.subheader("📥 Download dos Documentos")
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nome_zip = f"Documentos_Defesa_{st.session_state.periodo_letivo}_{timestamp}.zip"
    
    st.download_button(
        label="📦 Baixar Arquivo ZIP Completo",
        data=st.session_state.zip_buffer,
        file_name=nome_zip,
        mime="application/zip",
        use_container_width=True,
        help="Contém todas as declarações individuais + cronograma"
    )
    
    st.info("""
    **Conteúdo do arquivo ZIP:**
    - Declaração_Aluno1.docx
    - Declaração_Aluno2.docx
    - ...
    - Cronograma_defesas_monografia_[PERIODO].docx
    """)

# Mensagem inicial
if not uploaded_file:
    st.info("👈 **Para começar:** Clique em 'Browse files' e faça upload de um arquivo Excel acima.")
    
    st.markdown("---")
    st.subheader("🎯 Sobre o Sistema")
    st.markdown("""
    Este sistema gera automaticamente:
    
    1. **Declarações individuais** de banca examinadora para cada aluno
    2. **Cronograma completo** de defesas de monografia
    
    **Compatível com:**
    - Bacharelado em Química
    - Bacharelado em Química Industrial  
    - Licenciatura em Química
    
    **Funcionalidades:**
    - ✅ Formatação automática de nomes
    - ✅ Extração correta de horários e datas
    - ✅ Inclusão de todos os membros da banca
    - ✅ Cronograma em tabela única contínua
    - ✅ Data formatada sem dia da semana
    """)
